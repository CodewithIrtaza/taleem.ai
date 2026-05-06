from docx import Document as DocxDoc
from docx.shared import Pt, RGBColor
import os, uuid

def export_to_docx(title: str, transcript: str, notes: str) -> str:
    doc = DocxDoc()
    
    # Clean title (remove markdown or section markers if any)
    clean_title = title.replace('[', '').replace(']', '')
    doc.add_heading(f"Taleem.AI — {clean_title}", 0)

    doc.add_heading("Lecture Notes", 1)
    
    # Parse notes and apply formatting
    for line in notes.split("\n"):
        if line.strip().startswith('[') and line.strip().endswith(']'):
            # Section title - bold and black
            section_title = line.strip()[1:-1]  # Remove [ and ]
            p = doc.add_paragraph(section_title)
            p_format = p.style.font
            p_format.bold = True
            p_format.size = Pt(13)
            p_format.color.rgb = RGBColor(0, 0, 0)  # Standard Black color
        elif line.strip():
            # Regular content
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)
        else:
            # Empty line - just add a small break
            doc.add_paragraph()

    doc.add_heading("Full Transcript", 1)
    doc.add_paragraph(transcript)

    path = f"uploads/notes_{uuid.uuid4()}.docx"
    doc.save(path)
    return path